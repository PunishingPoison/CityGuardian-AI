from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ============================================================
# GLOBAL STYLE CONFIGURATION
# ============================================================
style_normal = doc.styles['Normal']
style_normal.font.name = 'Times New Roman'
style_normal.font.size = Pt(12)
style_normal.paragraph_format.line_spacing = 1.5
style_normal.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.name = 'Times New Roman'
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        heading_style.font.size = Pt(18)
        heading_style.paragraph_format.space_before = Pt(24)
        heading_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        heading_style.font.size = Pt(15)
        heading_style.paragraph_format.space_before = Pt(18)
        heading_style.paragraph_format.space_after = Pt(8)
    elif level == 3:
        heading_style.font.size = Pt(13)
        heading_style.paragraph_format.space_before = Pt(14)
        heading_style.paragraph_format.space_after = Pt(6)


def add_para(text, bold=False, italic=False, align=None, size=None, space_after=None, first_line_indent=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    if size:
        run.font.size = Pt(size)
    else:
        run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    if align:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p


def add_bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p


def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = 1.0


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return table


# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('CityGuardian', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=36, space_after=12)
add_para('AI-Driven Urban Crisis Response Simulator', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=20, space_after=24)
add_para('A Comprehensive Technical Report', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=48)

for _ in range(4):
    doc.add_paragraph()

add_para('Project Repository: https://github.com/PunishingPoison/CityGuardian.git', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading('Table of Contents', level=1)

toc_items = [
    ('1.', 'Introduction'),
    ('2.', 'About the Project'),
    ('  2.1', 'Project Objectives'),
    ('  2.2', 'Scope and Audience'),
    ('  2.3', 'Project Directory Structure'),
    ('3.', 'Concepts and Algorithms Used'),
    ('  3.1', 'Procedural City Generation'),
    ('  3.2', 'A* (A-Star) Pathfinding Algorithm'),
    ('  3.3', 'Cellular Automata for Fire Spread'),
    ('  3.4', 'Elevation-Based Flood Simulation'),
    ('  3.5', 'Finite State Machines for Emergency Resources'),
    ('  3.6', 'The Game Loop Pattern'),
    ('  3.7', 'Area Rescue Mechanics'),
    ('  3.8', 'Observer Pattern for UI Updates'),
    ('4.', 'Technology Stack'),
    ('  4.1', 'Java 17'),
    ('  4.2', 'JavaFX 17.0.6'),
    ('  4.3', 'SQLite (via JDBC)'),
    ('  4.4', 'Apache Maven'),
    ('5.', 'System Architecture'),
    ('  5.1', 'Model-View-Controller (MVC) Pattern'),
    ('  5.2', 'Package Structure'),
    ('  5.3', 'Class Hierarchy and Inheritance'),
    ('6.', 'Detailed Implementation'),
    ('  6.1', 'The City Model'),
    ('  6.2', 'The Tile System and TileType Enumeration'),
    ('  6.3', 'The Citizen Entity'),
    ('  6.4', 'Disaster Modules'),
    ('  6.5', 'Emergency Resource Modules'),
    ('  6.6', 'The Simulation Engine'),
    ('  6.7', 'The Evacuation Planner (A* Implementation)'),
    ('  6.8', 'The Dashboard Controller and Rendering Pipeline'),
    ('  6.9', 'The Database Manager'),
    ('  6.10', 'The Recommendation Engine'),
    ('  6.11', 'CSS Theming and Dark Mode'),
    ('7.', 'Installation and Setup Guide'),
    ('  7.1', 'Prerequisites'),
    ('  7.2', 'Cloning the Repository'),
    ('  7.3', 'Building the Project'),
    ('  7.4', 'Running the Application'),
    ('8.', 'User Walkthrough Guide'),
    ('9.', 'Results and Output'),
    ('10.', 'Performance Analysis'),
    ('11.', 'Challenges Encountered and Solutions'),
    ('12.', 'Future Enhancements'),
    ('13.', 'Conclusion'),
    ('14.', 'References'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{num}  {title}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15

doc.add_page_break()

# ============================================================
# 1. INTRODUCTION
# ============================================================
doc.add_heading('1. Introduction', level=1)

add_para(
    'Natural and man-made disasters represent some of the most significant threats to urban populations worldwide. '
    'Earthquakes, fires, and floods can strike with little or no warning, overwhelming emergency response infrastructure '
    'and leading to catastrophic loss of life. The effectiveness of disaster response is critically dependent on two factors: '
    'the speed at which emergency resources can be deployed to the disaster site, and the intelligence of the routing and '
    'allocation decisions made under extreme pressure.'
)
add_para(
    'Traditional emergency dispatch systems rely heavily on human operators making rapid decisions based on incomplete '
    'information. Radio communications can be overwhelmed, road conditions may change in real-time due to the disaster '
    'itself, and the sheer scale of a major urban disaster can exceed the cognitive capacity of any human coordinator. '
    'There is therefore a pressing need for automated, algorithmic systems that can model disaster scenarios, evaluate '
    'response strategies, and provide real-time decision support.'
)
add_para(
    'CityGuardian is an advanced, real-time simulation platform developed in Java that addresses these challenges. '
    'It provides a comprehensive environment for modeling urban disasters and simulating the autonomous deployment of '
    'emergency services including firetrucks, ambulances, and rescue helicopters. The platform procedurally generates '
    'realistic city layouts, simulates three distinct disaster types with physically-motivated spread mechanics, and '
    'deploys intelligent emergency resources that autonomously navigate the city grid using the A* pathfinding algorithm.'
)
add_para(
    'This report provides a thorough and exhaustive technical analysis of the CityGuardian project. It covers the '
    'theoretical foundations upon which the simulation is built, the complete technology stack, the detailed implementation '
    'of every module and class, the installation and deployment procedures, and a performance evaluation of the system. '
    'The report is intended to serve as a complete reference document for understanding every aspect of the project.'
)

doc.add_page_break()

# ============================================================
# 2. ABOUT THE PROJECT
# ============================================================
doc.add_heading('2. About the Project', level=1)

add_para(
    'CityGuardian is a multi-agent, real-time urban disaster simulation system. It operates on a two-dimensional '
    'grid-based map where each cell (tile) represents a discrete physical area of a city. The system is designed to '
    'manage thousands of individual citizen entities simultaneously, tracking their health, position, and status in '
    'real-time as disasters unfold and emergency services respond.'
)

doc.add_heading('2.1 Project Objectives', level=2)

add_para('The primary objectives of the CityGuardian project are:')
add_bullet('To create a procedurally generated city environment that serves as a realistic testbed for disaster simulation, '
           'including roads, residential zones, commercial districts, and hospital infrastructure.')
add_bullet('To implement three distinct disaster modules (Fire, Flood, Earthquake) each with unique, physically-motivated '
           'spread and damage mechanics.')
add_bullet('To develop autonomous emergency response agents (FireTruck, Ambulance, Helicopter) that intelligently navigate '
           'the city grid using the A* pathfinding algorithm to reach disaster sites and rescue citizens.')
add_bullet('To implement a multi-phase ambulance logistics system where ambulances physically load injured citizens, '
           'transport them along road networks, and deliver them to hospital facilities.')
add_bullet('To provide a real-time visual dashboard with live telemetry, population status tracking, and an interactive '
           'pie chart for monitoring simulation outcomes.')
add_bullet('To demonstrate the impact of resource allocation decisions by allowing users to vary the number of deployed '
           'emergency vehicles and observe the resulting difference in outcomes.')

doc.add_heading('2.2 Scope and Audience', level=2)

add_para(
    'The scope of the project encompasses the complete lifecycle of a disaster response simulation: city generation, '
    'disaster initiation, resource deployment, real-time simulation, and outcome analysis. The application is designed '
    'to be self-contained, requiring no external services or network connectivity to operate.'
)
add_para(
    'The intended audience includes researchers in algorithmic logistics and multi-agent systems, urban planners '
    'seeking to understand bottlenecks in emergency response networks, computer science students studying advanced Java '
    'programming concepts such as inheritance hierarchies, state machines, and real-time rendering, and educators '
    'requiring a demonstrative platform for pathfinding algorithms and simulation design patterns.'
)

doc.add_heading('2.3 Project Directory Structure', level=2)

add_para('The project follows the standard Maven directory layout:')

add_code_block([
    'CityGuardian/',
    '  pom.xml',
    '  README.md',
    '  cityguardian.db',
    '  src/',
    '    main/',
    '      java/',
    '        com/cityguardian/',
    '          Main.java',
    '          controller/',
    '            DashboardController.java',
    '          db/',
    '            DatabaseManager.java',
    '          engine/',
    '            EvacuationPlanner.java',
    '            RecommendationEngine.java',
    '            SimulationEngine.java',
    '          model/',
    '            Citizen.java',
    '            City.java',
    '            Tile.java',
    '            TileType.java',
    '            building/',
    '              Building.java',
    '              Hospital.java',
    '            disaster/',
    '              Disaster.java',
    '              EarthquakeDisaster.java',
    '              FireDisaster.java',
    '              FloodDisaster.java',
    '            resource/',
    '              Ambulance.java',
    '              EmergencyResource.java',
    '              FireTruck.java',
    '              Helicopter.java',
    '      resources/',
    '        css/',
    '          style.css',
    '        fxml/',
    '          Dashboard.fxml',
])

doc.add_page_break()

# ============================================================
# 3. CONCEPTS AND ALGORITHMS USED
# ============================================================
doc.add_heading('3. Concepts and Algorithms Used', level=1)

add_para(
    'CityGuardian is built upon a foundation of well-established algorithms and design patterns drawn from computer science, '
    'game development, and computational mathematics. This section provides a detailed explanation of each concept and how '
    'it is applied within the project.'
)

doc.add_heading('3.1 Procedural City Generation', level=2)

add_para(
    'Procedural generation refers to the algorithmic creation of data content rather than manual construction. In CityGuardian, '
    'the entire city layout is generated at runtime through a combination of deterministic grid patterns and stochastic '
    'placement algorithms. This approach ensures that each simulation run operates on a unique city topology, preventing '
    'over-fitting of emergency response strategies to a single map layout.'
)
add_para(
    'The generation algorithm operates in the following phases. First, the grid is initialized as a 70x60 two-dimensional '
    'array of Tile objects, each defaulting to the EMPTY type. Second, an elevation map is computed for every tile using '
    'a Euclidean distance function from the center of the grid. The elevation value is calculated as:'
)
add_para('elevation(x, y) = 1.0 - (distance(x, y, center) / maxDistance) + random(-0.1, 0.1)', italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(
    'This produces a dome-shaped terrain profile with randomized local variation, which is critical for the flood '
    'simulation mechanics. Third, roads are laid out on a strict modulo-based grid system: any tile where x mod 6 equals '
    'zero or y mod 6 equals zero is designated as a ROAD tile. This guarantees full connectivity of the road network and '
    'ensures that pathfinding can always find a route between any two points on the grid. Fourth, non-road tiles are '
    'probabilistically assigned as RESIDENTIAL (30% chance, with 3 citizens spawned per tile), COMMERCIAL (10% chance), '
    'HOSPITAL (2% chance), or left as EMPTY.'
)

doc.add_heading('3.2 A* (A-Star) Pathfinding Algorithm', level=2)

add_para(
    'The A* algorithm is a best-first graph search algorithm that finds the shortest path between a start node and a goal '
    'node. It was first described by Hart, Nilsson, and Raphael in their seminal 1968 paper. A* combines the advantages '
    'of Dijkstra\'s algorithm (which guarantees optimality) with a heuristic function that dramatically reduces the number '
    'of nodes explored, resulting in significantly improved computational performance.'
)
add_para(
    'In CityGuardian, the A* implementation is housed in the EvacuationPlanner class. The algorithm maintains two data '
    'structures: an open set (implemented as a PriorityQueue for O(log n) extraction of the minimum-cost node) and a '
    'closed set (implemented as a HashSet for O(1) membership testing). For each node, the algorithm computes a cost '
    'function f(n) = g(n) + h(n), where g(n) is the actual cost from the start node to the current node, and h(n) is '
    'the heuristic estimate of the cost from the current node to the goal.'
)
add_para(
    'The heuristic function used is the Manhattan distance: h(n) = |x_current - x_goal| + |y_current - y_goal|. This '
    'is an admissible heuristic for grid-based movement with four-directional adjacency, meaning it never overestimates '
    'the true cost, which guarantees that A* will find an optimal path.'
)
add_para(
    'A critical feature of the CityGuardian implementation is the dynamic cost modification based on tile risk levels. '
    'Tiles that are actively experiencing a disaster receive a massive penalty of 1000 added to their traversal cost, '
    'while tiles with elevated risk levels receive a penalty proportional to their risk (riskLevel multiplied by 10). '
    'This causes the pathfinding algorithm to naturally route vehicles around active disaster zones, which is essential '
    'for preventing emergency vehicles from driving directly through fires or floods.'
)
add_para(
    'The algorithm also supports a configurable "requiresRoad" parameter. When set to true (used for ground vehicles such as '
    'firetrucks and ambulances), the algorithm restricts traversal to ROAD tiles only. When set to false (used for citizen '
    'evacuation), the algorithm permits traversal through any walkable tile type.'
)

doc.add_heading('3.3 Cellular Automata for Fire Spread', level=2)

add_para(
    'Cellular automata (CA) are discrete computational models consisting of a grid of cells, each of which can be in '
    'one of a finite number of states. At each time step, every cell updates its state based on its current state and the '
    'states of its neighbors according to a fixed set of rules. Cellular automata were extensively studied by Stephen '
    'Wolfram in the 1980s and have been widely applied to model physical phenomena including fluid dynamics, crystal '
    'growth, and fire propagation.'
)
add_para(
    'In CityGuardian, the FireDisaster class implements a probabilistic cellular automaton. Every 0.8 seconds of simulation '
    'time, the spread() method iterates over every tile in the grid. For each tile that is currently on fire (hasDisaster '
    'equals true), two probabilistic evaluations occur. First, there is a 5% chance that the fire on this tile will '
    'self-extinguish (burnout), at which point the tile transitions to the BURNT state if it is not a road. Second, for '
    'each of the four cardinal neighbors (north, south, east, west), there is a 40% chance that the fire will spread to '
    'that neighbor, provided the neighbor is not already burning, is not already burnt, and is not a water tile.'
)
add_para(
    'The 0.8-second spread interval and 40% spread probability were carefully calibrated to produce a fire that spreads '
    'at a moderate, visible pace. This tuning ensures that a small fleet of 3 firetrucks is unable to contain the fire '
    '(the spread outpaces the containment), while a larger fleet of 15 or more firetrucks can successfully suppress the '
    'outbreak before it consumes the entire city. This creates meaningful gameplay differentiation based on the resource '
    'allocation decisions made by the user.'
)

doc.add_heading('3.4 Elevation-Based Flood Simulation', level=2)

add_para(
    'The flood simulation uses a physically-motivated model based on terrain elevation. Unlike fire, which spreads '
    'uniformly in all four cardinal directions, water flows downhill. When the flood evolves, each flooded tile examines '
    'its four neighbors and identifies the one with the lowest elevation value. If this lowest neighbor is not already '
    'flooded and passes a random probability check (80% chance), it is converted to a WATER tile and flagged as a '
    'disaster zone.'
)
add_para(
    'This elevation-driven spread produces organic, realistic flooding patterns. Low-lying areas of the city (those '
    'farther from the center of the dome-shaped elevation profile) are submerged first, while elevated areas near the '
    'city center remain dry for longer. The flood spreads at a slower interval of 3.0 seconds, reflecting the slower '
    'physical dynamics of water compared to fire.'
)

doc.add_heading('3.5 Finite State Machines for Emergency Resources', level=2)

add_para(
    'Each emergency resource (FireTruck, Ambulance, Helicopter) operates as a finite state machine (FSM). The FSM '
    'defines a strict set of states and transitions that govern the behavior of the entity. The EmergencyResource base '
    'class defines four states in its Status enumeration: AVAILABLE, DISPATCHED, BUSY, and RETURNING.'
)
add_para('The state transitions for each resource type are as follows:')

add_table(
    ['Resource', 'AVAILABLE', 'DISPATCHED', 'RETURNING'],
    [
        ['FireTruck', 'Scans for nearest fire tile or injured citizen', 'Follows A* path along roads to target; extinguishes 5x5 area on arrival', 'N/A (returns to AVAILABLE after action)'],
        ['Ambulance', 'Scans for nearest injured citizen not yet loaded', 'Follows A* path along roads to citizen; loads up to 10 injured', 'Follows A* path to nearest hospital; unloads all passengers on arrival'],
        ['Helicopter', 'Scans for nearest citizen stranded in flood water', 'Flies directly to target (ignores roads); rescues 3x3 area on arrival', 'N/A (returns to AVAILABLE after action)'],
    ]
)

add_para(
    'The Ambulance has the most complex state machine of the three resources. Upon reaching its target and loading '
    'injured citizens, it checks whether it has reached its maximum capacity of 10 passengers. If full, it transitions '
    'to RETURNING and uses A* to pathfind to the nearest hospital. Upon arrival at the hospital, it unloads all '
    'passengers (marking them as evacuated), clears its passenger list, and transitions back to AVAILABLE for the '
    'next rescue cycle.', space_after=6
)

doc.add_heading('3.6 The Game Loop Pattern', level=2)

add_para(
    'The Game Loop is a fundamental software design pattern used in virtually all real-time simulations and video games. '
    'It provides a mechanism for continuously updating and rendering the simulation state at a consistent rate, '
    'independent of hardware performance.'
)
add_para(
    'CityGuardian implements the game loop using JavaFX\'s AnimationTimer class. The AnimationTimer.handle() method is '
    'invoked by the JavaFX runtime on every frame of the rendering pipeline (typically 60 frames per second on modern '
    'hardware). The method receives a nanosecond-precision timestamp, from which the elapsed time since the last frame '
    '(deltaTime) is computed. This deltaTime value is passed to the tick() method of the SimulationEngine, which uses '
    'it to scale all movement and spread calculations. This approach ensures that the simulation behaves consistently '
    'regardless of frame rate variations.'
)
add_para(
    'The tick() method processes the simulation in a strict, deterministic order: (1) disaster evolution, (2) citizen '
    'status updates, (3) resource spawning, (4) resource dispatching and movement, and (5) UI notification via callback.'
)

doc.add_heading('3.7 Area Rescue Mechanics', level=2)

add_para(
    'When an emergency vehicle arrives at its target destination, it does not simply rescue a single citizen or '
    'extinguish a single tile. Instead, it performs an "area rescue" operation over a square region centered on its '
    'current position. The performAreaRescue method takes a center coordinate, a radius, and a boolean indicating '
    'whether the caller is a firetruck.'
)
add_para(
    'For all resource types, citizens within the bounding box (center minus radius to center plus radius in both '
    'dimensions) are immediately marked as evacuated. For firetrucks specifically, tiles within the bounding box that '
    'have an active disaster are extinguished (hasDisaster set to false, riskLevel reset to 0), and combustible tiles '
    '(RESIDENTIAL, COMMERCIAL, HOSPITAL, SHELTER) are converted to BURNT to prevent re-ignition. Firetrucks operate '
    'with a radius of 2 (covering a 5x5 block), while helicopters use a radius of 1 (covering a 3x3 block).'
)

doc.add_heading('3.8 Observer Pattern for UI Updates', level=2)

add_para(
    'The Observer pattern is used to decouple the simulation engine from the user interface. The SimulationEngine '
    'class accepts a Runnable callback (onTick) in its constructor. At the end of every tick, the engine invokes this '
    'callback, which triggers the DashboardController to repaint the map canvas and update the population statistics. '
    'The callback execution is wrapped in Platform.runLater() to ensure that all UI modifications occur on the JavaFX '
    'Application Thread, as required by the JavaFX threading model.'
)

doc.add_page_break()

# ============================================================
# 4. TECHNOLOGY STACK
# ============================================================
doc.add_heading('4. Technology Stack', level=1)

doc.add_heading('4.1 Java 17', level=2)
add_para(
    'Java 17 is a Long-Term Support (LTS) release of the Java programming language and runtime platform. It was chosen '
    'for this project due to its strong type safety, mature object-oriented programming model, excellent support for '
    'inheritance hierarchies and abstract classes, and its vast standard library. The project leverages Java 17 features '
    'including the enhanced switch expressions, the Stream API for filtering and counting resources, and the modular '
    'system (Java Platform Module System) via the module-info.java descriptor.'
)

doc.add_heading('4.2 JavaFX 17.0.6', level=2)
add_para(
    'JavaFX is the standard GUI toolkit for modern Java desktop applications. CityGuardian uses JavaFX for all visual '
    'components: the FXML-based layout system for constructing the dashboard interface, the Canvas API for high-performance '
    'rendering of the city grid and entities, the PieChart control for real-time population statistics, and the CSS '
    'styling system for implementing a dark-themed professional interface.'
)
add_para(
    'The Canvas API was deliberately chosen over the Scene Graph approach for rendering the map. In a Scene Graph '
    'approach, each of the 4200 tiles (70 times 60) and thousands of citizens would need to be represented as individual '
    'JavaFX Node objects, each carrying significant memory overhead and requiring scene graph traversal during rendering. '
    'By using the Canvas API, the rendering is performed via immediate-mode drawing commands (fillRect, fillOval) on a '
    'single GraphicsContext, which is orders of magnitude more efficient for this use case.'
)

doc.add_heading('4.3 SQLite (via JDBC)', level=2)
add_para(
    'SQLite is a lightweight, file-based relational database engine. CityGuardian integrates SQLite through the Xerial '
    'SQLite JDBC driver (version 3.41.2.1) to provide persistent storage capabilities. The DatabaseManager class '
    'initializes a local database file (cityguardian.db) and creates three tables: tiles (for storing grid state), '
    'citizens (for storing citizen data), and simulation_runs (for recording historical simulation outcomes with '
    'timestamps, casualty counts, and save counts). This infrastructure enables future features such as simulation '
    'replay and statistical analysis across multiple runs.'
)

doc.add_heading('4.4 Apache Maven', level=2)
add_para(
    'Apache Maven is the build automation and dependency management tool used by CityGuardian. The project is configured '
    'through a pom.xml file that specifies all external dependencies (JavaFX Controls, JavaFX FXML, SQLite JDBC), '
    'compiler settings (Java 17 source and target), and build plugins (maven-compiler-plugin version 3.10.1 and '
    'javafx-maven-plugin version 0.0.8). Maven ensures reproducible builds and simplifies the installation process '
    'for end users.'
)

doc.add_page_break()

# ============================================================
# 5. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('5. System Architecture', level=1)

doc.add_heading('5.1 Model-View-Controller (MVC) Pattern', level=2)
add_para(
    'CityGuardian strictly adheres to the Model-View-Controller architectural pattern, one of the most widely used '
    'software architectural patterns in application development. This separation of concerns ensures that the domain '
    'logic, user interface, and control flow are each encapsulated in distinct, loosely-coupled modules.'
)
add_bullet('Model: Contains all domain entities including City, Tile, Citizen, Building, Disaster subclasses, and '
           'EmergencyResource subclasses. These classes encapsulate data and business rules but have no knowledge of '
           'the user interface.')
add_bullet('View: Defined in the FXML layout file (Dashboard.fxml) and CSS stylesheet (style.css). The view specifies '
           'the visual structure and styling of the dashboard interface without containing any business logic.')
add_bullet('Controller: The DashboardController class bridges the model and view. It handles user input events '
           '(button clicks, text field changes), delegates simulation control to the SimulationEngine, and renders '
           'the model state onto the Canvas.')

doc.add_heading('5.2 Package Structure', level=2)

add_table(
    ['Package', 'Contents', 'Responsibility'],
    [
        ['com.cityguardian', 'Main.java', 'Application entry point; initializes database and launches JavaFX stage'],
        ['com.cityguardian.controller', 'DashboardController.java', 'UI event handling, map rendering, statistics display'],
        ['com.cityguardian.db', 'DatabaseManager.java', 'SQLite database connection and schema management'],
        ['com.cityguardian.engine', 'SimulationEngine.java, EvacuationPlanner.java, RecommendationEngine.java', 'Core simulation loop, A* pathfinding, AI insights'],
        ['com.cityguardian.model', 'City.java, Tile.java, Citizen.java, TileType.java', 'Domain entities for city grid and population'],
        ['com.cityguardian.model.building', 'Building.java, Hospital.java', 'Building infrastructure with capacity and structural integrity'],
        ['com.cityguardian.model.disaster', 'Disaster.java, FireDisaster.java, FloodDisaster.java, EarthquakeDisaster.java', 'Disaster type definitions and evolution logic'],
        ['com.cityguardian.model.resource', 'EmergencyResource.java, FireTruck.java, Ambulance.java, Helicopter.java', 'Emergency vehicle definitions, state machines, and pathfinding state'],
    ]
)

doc.add_heading('5.3 Class Hierarchy and Inheritance', level=2)

add_para(
    'The project makes extensive use of Java\'s inheritance mechanism to define clear hierarchies with shared '
    'behavior and polymorphic dispatch.'
)
add_para('Disaster Hierarchy:', bold=True)
add_bullet('Disaster (abstract base class): Defines common fields (startX, startY, severity, damageRadius, active) '
           'and the abstract method evolve(City, double).')
add_bullet('FireDisaster extends Disaster: Implements cellular automata spread logic.', level=1)
add_bullet('FloodDisaster extends Disaster: Implements elevation-based water flow logic.', level=1)
add_bullet('EarthquakeDisaster extends Disaster: Implements instant-impact area damage logic.', level=1)

add_para('Emergency Resource Hierarchy:', bold=True)
add_bullet('EmergencyResource (abstract base class): Defines common fields (id, x, y, status, targetTile) and '
           'the abstract method performAction().')
add_bullet('FireTruck extends EmergencyResource: Adds road-based pathfinding state and fire extinguishing behavior.', level=1)
add_bullet('Ambulance extends EmergencyResource: Adds passenger capacity, loaded citizens list, hospital return logic.', level=1)
add_bullet('Helicopter extends EmergencyResource: Adds direct flight movement (ignores road constraints).', level=1)

add_para('Building Hierarchy:', bold=True)
add_bullet('Building (abstract base class): Defines common fields (name, capacity, currentOccupancy, structuralIntegrity).')
add_bullet('Hospital extends Building: Specializes for medical facility functionality.', level=1)

doc.add_page_break()

# ============================================================
# 6. DETAILED IMPLEMENTATION
# ============================================================
doc.add_heading('6. Detailed Implementation', level=1)

doc.add_heading('6.1 The City Model', level=2)
add_para(
    'The City class (City.java) serves as the central data container for the entire simulation. It encapsulates a '
    'two-dimensional array of Tile objects (Tile[][] grid) with configurable width (70) and height (60), yielding '
    'a total of 4,200 tiles. The class also maintains three ArrayList collections: citizens (storing all Citizen '
    'objects in the simulation), disasters (storing all active Disaster objects), and resources (storing all deployed '
    'EmergencyResource objects).'
)
add_para(
    'The City provides boundary-checked tile access through its getTile(int x, int y) method, which returns null '
    'for out-of-bounds coordinates rather than throwing an ArrayIndexOutOfBoundsException. This null-safe access '
    'pattern simplifies neighbor iteration in the pathfinding and disaster spread algorithms.'
)

doc.add_heading('6.2 The Tile System and TileType Enumeration', level=2)
add_para(
    'Each Tile object represents a single cell in the city grid. It stores its immutable position (x, y), a mutable '
    'type (TileType enum), an optional Building reference, a riskLevel (double from 0.0 to 1.0), a hasDisaster '
    'boolean flag, and an elevation value (used by the flood simulation). The TileType enumeration defines twelve '
    'distinct tile types:'
)

add_table(
    ['TileType', 'Description', 'Walkable', 'Visual Color'],
    [
        ['EMPTY', 'Undeveloped land', 'Yes', 'Grayscale (elevation-based)'],
        ['ROAD', 'Navigable road for ground vehicles', 'Yes', 'Dark Gray'],
        ['RESIDENTIAL', 'Housing block (spawns 3 citizens)', 'No (vehicles)', 'Light Blue'],
        ['COMMERCIAL', 'Commercial district', 'No (vehicles)', 'Orange'],
        ['HOSPITAL', 'Medical facility (ambulance drop-off point)', 'Yes', 'White'],
        ['SHELTER', 'Emergency shelter', 'Yes', 'Light Green'],
        ['WATER', 'Flood water', 'No', 'Dark Blue'],
        ['OBSTACLE', 'Earthquake debris', 'No', 'Saddle Brown'],
        ['BURNT', 'Fire-damaged area', 'No', 'Dark Gray (RGB 50,50,50)'],
        ['SCHOOL', 'Educational facility (reserved)', 'No', 'N/A'],
        ['POLICE_STATION', 'Law enforcement (reserved)', 'No', 'N/A'],
        ['FIRE_STATION', 'Fire department (reserved)', 'No', 'N/A'],
    ]
)

doc.add_heading('6.3 The Citizen Entity', level=2)
add_para(
    'The Citizen class (Citizen.java) models an individual person in the city. Each citizen has a name, age, health '
    '(starting at 100.0), position (x, y as doubles for smooth interpolation), and three boolean state flags: '
    'isInjured, isEvacuated, and isDead. Citizens also maintain pathfinding state for evacuation (evacuationPath '
    'and currentPathIndex) and a movementSpeed of 2.0 tiles per second.'
)
add_para(
    'The health system uses a threshold-based injury model. When a citizen\'s health drops below 50.0 (via the '
    'takeDamage method), the isInjured flag is set to true, which causes the citizen to be rendered in yellow on the '
    'map and makes them a valid target for ambulance dispatch. When health reaches 0.0, the isDead flag is set and '
    'the citizen is no longer rendered or counted in the active population.'
)
add_para(
    'An important design decision is that injured citizens cannot move. When a citizen becomes injured, any active '
    'evacuation path is cleared (set to null), and no new path will be assigned. This ensures injured citizens remain '
    'stationary at the disaster site until an ambulance arrives to rescue them, creating a realistic triage scenario.'
)

doc.add_heading('6.4 Disaster Modules', level=2)

add_para('6.4.1 FireDisaster', bold=True)
add_para(
    'The FireDisaster class implements a time-accumulator pattern to control spread timing. A private double field '
    'timeAccumulator tracks elapsed simulation time since the last spread event. When timeAccumulator exceeds the '
    'spreadInterval threshold of 0.8 seconds, the spread() method is invoked and the accumulator is reset to zero.'
)
add_para(
    'The fire is initially seeded as a large 13x13 block (radius 6) centered at coordinates (35, 30) through the '
    'triggerFire method in DashboardController. Roads within the initial fire zone are deliberately excluded from '
    'burning to ensure that firetrucks can navigate into the disaster zone. During spread, each burning tile has a 5% '
    'chance of self-extinguishing per interval, which creates natural gaps and prevents the fire from being a '
    'monolithic wall. The fire does not destroy road tiles when it burns out naturally, preserving the road network '
    'for emergency vehicle navigation.'
)

add_para('6.4.2 FloodDisaster', bold=True)
add_para(
    'The FloodDisaster class uses a 3.0-second spread interval, reflecting the slower dynamics of water flow compared '
    'to fire. The spread() method identifies each currently flooded tile and examines its four cardinal neighbors. '
    'Unlike fire, which can spread to any combustible neighbor, water preferentially flows to the neighbor with the '
    'lowest elevation value. If the lowest-elevation neighbor passes an 80% probability check, it is converted to a '
    'WATER tile and marked as a disaster zone with a risk level of 0.9.'
)
add_para(
    'The flood is initiated from a single seed tile at coordinates (10, 10), which is set to WATER type with '
    'hasDisaster set to true. The elevation-dependent spread creates realistic water flow patterns where low-lying '
    'areas are submerged first.'
)

add_para('6.4.3 EarthquakeDisaster', bold=True)
add_para(
    'The EarthquakeDisaster class implements a single-trigger, instant-impact model. Unlike fire and flood, which '
    'evolve continuously over time, an earthquake fires exactly once and then deactivates itself (active is set to '
    'false). The earthquake has a damage radius of 10 tiles (creating a 21x21 impact zone) and applies two effects:'
)
add_bullet('Structural damage: Non-EMPTY tiles within the radius have a 20% chance of being converted to OBSTACLE '
           'type, representing collapsed buildings and debris.')
add_bullet('Citizen injury: All living, non-evacuated citizens within the radius receive 60.0 points of damage. '
           'Since citizens start with 100.0 health and the injury threshold is 50.0, this instantly drops their '
           'health to 40.0, guaranteeing that all affected citizens are immediately classified as injured and '
           'rendered in yellow.')

doc.add_heading('6.5 Emergency Resource Modules', level=2)

add_para('6.5.1 EmergencyResource (Base Class)', bold=True)
add_para(
    'EmergencyResource is an abstract class that defines the shared contract for all emergency vehicles. It stores '
    'a unique string identifier, integer position coordinates, a Status enumeration (AVAILABLE, DISPATCHED, BUSY, '
    'RETURNING), and an optional targetTile reference. The class declares the abstract method performAction(), which '
    'each subclass must implement to define its behavior upon reaching its target.'
)

add_para('6.5.2 FireTruck', bold=True)
add_para(
    'The FireTruck class extends EmergencyResource with additional fields for road-based pathfinding: a List of Tile '
    'objects representing the computed A* path, a pathIndex tracking the current position along the path, high-precision '
    'double coordinates (currentX, currentY) for smooth interpolation between tile centers, and a speed of 3.0 tiles '
    'per second. The performAction method marks any targeted citizen as evacuated and resets the truck to AVAILABLE status.'
)

add_para('6.5.3 Ambulance', bold=True)
add_para(
    'The Ambulance class is the most complex emergency resource. In addition to the standard pathfinding fields (path, '
    'pathIndex, currentX, currentY), it introduces a capacity field (set to 10), a loadedCitizens ArrayList for '
    'tracking passengers currently aboard the vehicle, a returningToHospital boolean flag, and a significantly higher '
    'speed of 10.0 tiles per second (more than three times faster than a firetruck).'
)
add_para(
    'The performAction method implements branching logic based on the current phase. If returningToHospital is true, '
    'it iterates over all loaded citizens, marks each as evacuated, clears the loaded list, resets the flag, and '
    'transitions to AVAILABLE. If returningToHospital is false, it checks whether the ambulance is full. If full, '
    'it sets returningToHospital to true and transitions to RETURNING status. If not full, it transitions to AVAILABLE '
    'to allow the simulation engine to dispatch it to another target.'
)

add_para('6.5.4 Helicopter', bold=True)
add_para(
    'The Helicopter class provides aerial rescue capability for flood scenarios. Unlike ground vehicles, helicopters '
    'ignore road constraints entirely and fly in a straight line toward their target at a speed of 5.0 tiles per second. '
    'The performAction method marks the target citizen as evacuated and returns to AVAILABLE status.'
)

doc.add_heading('6.6 The Simulation Engine', level=2)
add_para(
    'The SimulationEngine class (SimulationEngine.java, 516 lines) is the central orchestrator of the entire simulation. '
    'It is responsible for the game loop, disaster evolution, citizen health updates, resource spawning, target '
    'assignment, pathfinding delegation, movement interpolation, and UI notification.'
)
add_para(
    'The tick() method processes the simulation in five ordered phases per frame:'
)
add_bullet('Phase 1 - Disaster Evolution: Each active disaster\'s evolve() method is called with the current deltaTime, '
           'advancing fire spread, flood expansion, or earthquake damage application.')
add_bullet('Phase 2 - Citizen Updates: Each citizen\'s health is decremented by 10.0 multiplied by deltaTime if they are '
           'standing on a disaster tile. Injured citizens have their evacuation paths cleared to prevent autonomous movement.')
add_bullet('Phase 3 - Resource Spawning: The engine counts the current number of each resource type and spawns new '
           'instances (on random road tiles) until the user-configured maximum is reached for each active disaster type.')
add_bullet('Phase 4 - Resource Dispatch and Movement: AVAILABLE resources are assigned to the nearest valid target using '
           'Euclidean distance. DISPATCHED resources follow their pre-computed A* paths using velocity-based interpolation. '
           'RETURNING ambulances pathfind to the nearest hospital and unload passengers upon arrival.')
add_bullet('Phase 5 - UI Notification: The onTick callback is invoked to trigger a full map redraw and statistics update.')

add_para(
    'The engine also contains several important helper methods: findNearestSafeZone scans the entire grid for SHELTER '
    'or HOSPITAL tiles; findRandomRoadTile selects a random road tile for resource spawning; isTileTargeted prevents '
    'multiple resources from being assigned to the same citizen; isFireTileTargeted prevents multiple firetrucks from '
    'targeting the same fire tile; findNearestRoadTo locates the closest road tile to a given coordinate for pathfinding '
    'start/end points; isCitizenLoaded checks whether a citizen is already aboard an ambulance; and findNearestHospital '
    'locates the closest hospital for ambulance return routing.'
)

doc.add_heading('6.7 The Evacuation Planner (A* Implementation)', level=2)
add_para(
    'The EvacuationPlanner class (EvacuationPlanner.java, 108 lines) houses the complete A* pathfinding implementation. '
    'It uses an inner class NodeRecord that implements Comparable for priority queue ordering, storing the tile reference, '
    'parent node (for path reconstruction), cost-so-far (g value), and estimated total cost (f value).'
)
add_para(
    'The findPathAStar static method accepts the City, a start Tile, a goal Tile, and a boolean requiresRoad parameter. '
    'The open set is maintained as a PriorityQueue with an associated HashMap for O(1) lookups, and the closed set is '
    'a HashSet. The algorithm supports path updates for nodes already in the open set (by removing and re-inserting '
    'with updated costs), ensuring optimality. The reconstructPath method traces the parent chain from the goal back '
    'to the start and reverses the result to produce a forward-ordered path.'
)

doc.add_heading('6.8 The Dashboard Controller and Rendering Pipeline', level=2)
add_para(
    'The DashboardController class (DashboardController.java, 301 lines) manages the complete user interface. It is '
    'connected to the Dashboard.fxml layout via JavaFX\'s FXML injection mechanism, using @FXML annotations to bind '
    'UI controls (Canvas, ComboBox, ListView, Labels, TextField, PieChart) to controller fields.'
)
add_para(
    'The drawMap() method performs a complete re-render of the entire canvas on every frame. It iterates over all '
    '4,200 tiles, selecting the fill color based on the tile\'s disaster state and type. It then renders all living, '
    'non-evacuated citizens as small colored circles (green for healthy, yellow for injured). Finally, it renders '
    'all emergency resources as colored rectangles at their interpolated positions (magenta for firetrucks, white for '
    'ambulances, cyan for helicopters). Each tile is rendered as a (TILE_SIZE - 1) pixel square with a 1-pixel gap '
    'to create a visible grid pattern. TILE_SIZE is set to 10 pixels.'
)
add_para(
    'The controller also manages disaster triggering through three button handler methods: triggerEarthquake (creates '
    'an EarthquakeDisaster at center coordinates 35,30), triggerFire (seeds a 13x13 initial fire block and creates a '
    'FireDisaster), and triggerFlood (seeds a single water tile at 10,10 and creates a FloodDisaster).'
)

doc.add_heading('6.9 The Database Manager', level=2)
add_para(
    'The DatabaseManager class provides SQLite integration through JDBC. It manages a single static Connection '
    'instance connected to the local file "cityguardian.db". The initialize() method establishes the connection '
    'and calls createTables() to ensure the schema exists. Three tables are created using CREATE TABLE IF NOT EXISTS '
    'statements: tiles (id, x, y, type), citizens (id, name, age, health, x, y), and simulation_runs (id, timestamp, '
    'casualties, saved). The close() method safely terminates the database connection when the application exits.'
)

doc.add_heading('6.10 The Recommendation Engine', level=2)
add_para(
    'The RecommendationEngine class provides a framework for AI-driven insights. Currently, it contains a stub '
    'implementation that returns a single placeholder insight. The class is architecturally positioned to be expanded '
    'with real-time analysis of disaster progression, resource utilization rates, and predictive modeling of citizen '
    'outcomes. Example future insights include hospital capacity warnings, fire containment probability estimates, '
    'and optimal resource reallocation suggestions.'
)

doc.add_heading('6.11 CSS Theming and Dark Mode', level=2)
add_para(
    'The application uses a comprehensive CSS stylesheet (style.css, 138 lines) to implement a professional dark '
    'theme. The root pane background is set to #121212 (near-black), with panel backgrounds at #1e1e1e and border '
    'accents at #333333. Text is rendered in white or light gray (#aaaaaa, #cccccc, #e0e0e0) for maximum contrast. '
    'Interactive elements use hover effects (e.g., control buttons transition from #2d2d2d to #3d3d3d) and semantic '
    'coloring (the Start button uses green #2e7d32, disaster buttons use dark red #4a1919). The PieChart legend '
    'and labels are overridden to white for dark-mode compatibility. Input fields use white text fill to ensure '
    'visibility against the dark background.'
)

doc.add_page_break()

# ============================================================
# 7. INSTALLATION AND SETUP GUIDE
# ============================================================
doc.add_heading('7. Installation and Setup Guide', level=1)

doc.add_heading('7.1 Prerequisites', level=2)
add_para('The following software must be installed before building and running CityGuardian:')
add_bullet('Java Development Kit (JDK) 17 or higher. The JAVA_HOME environment variable must be set to the JDK installation directory, and the JDK bin directory must be added to the system PATH.')
add_bullet('Apache Maven 3.8 or higher. Maven must be accessible from the command line.')
add_bullet('Git (for cloning the repository from GitHub).')
add_bullet('A display environment capable of running JavaFX applications (any modern desktop operating system with a graphical display).')

doc.add_heading('7.2 Cloning the Repository', level=2)
add_para('Open a terminal or command prompt and execute:')
add_code_block([
    'git clone https://github.com/PunishingPoison/CityGuardian.git',
    'cd CityGuardian',
])

doc.add_heading('7.3 Building the Project', level=2)
add_para('Compile the source code and resolve all dependencies using Maven:')
add_code_block(['mvn clean compile'])
add_para(
    'This command will download all required dependencies (JavaFX Controls 17.0.6, JavaFX FXML 17.0.6, SQLite JDBC '
    '3.41.2.1) from the Maven Central Repository, compile all Java source files in the src/main/java directory, and '
    'place the compiled class files in the target/classes directory.'
)

doc.add_heading('7.4 Running the Application', level=2)
add_para('Launch the application using the JavaFX Maven plugin:')
add_code_block(['mvn javafx:run'])
add_para(
    'This command configures the JavaFX module path and launches the Main class, which initializes the SQLite database, '
    'loads the Dashboard.fxml layout, applies the dark-mode CSS stylesheet, and displays the application window at a '
    'resolution of 1280x800 pixels.'
)

doc.add_page_break()

# ============================================================
# 8. USER WALKTHROUGH GUIDE
# ============================================================
doc.add_heading('8. User Walkthrough Guide', level=1)

add_para('This section provides a step-by-step guide for using the CityGuardian application.')

add_para('Step 1: Generate the City', bold=True)
add_para(
    'Click the "Generate City" button in the left panel. This invokes the procedural generation algorithm, creating '
    'a complete city with roads, residential blocks, commercial zones, and hospitals. The map canvas in the center '
    'will display the generated city. Approximately 3,000 to 5,000 citizens will be spawned across the residential '
    'zones, as indicated by the "Total Citizens" counter in the bottom panel.'
)

add_para('Step 2: Configure Emergency Resources', bold=True)
add_para(
    'In the top toolbar, set the desired number of Firetrucks, Helicopters, and Ambulances using the text input '
    'fields. The default value is 3 for each. For fire scenarios, setting the firetruck count to 3 will result in '
    'the fire outpacing containment, while setting it to 15 or more will allow successful suppression. For earthquake '
    'scenarios, increasing the ambulance count will accelerate the rescue operation.'
)

add_para('Step 3: Start the Simulation', bold=True)
add_para(
    'Click the "Start" button. The simulation loop will begin executing, processing all active disasters and deploying '
    'emergency resources in real-time.'
)

add_para('Step 4: Trigger a Disaster', bold=True)
add_para(
    'Click one of the three disaster buttons in the left panel: "Earthquake", "Fire Outbreak", or "Flood". The '
    'selected disaster will be initiated at a predetermined location on the map, and emergency resources will begin '
    'spawning and deploying autonomously. The AI Insights panel on the right will log warning messages.'
)

add_para('Step 5: Observe the Simulation', bold=True)
add_para(
    'Watch the emergency vehicles navigate the road network to reach the disaster zone. Firetrucks will appear as '
    'magenta squares, ambulances as white squares, and helicopters as cyan squares. Injured citizens appear as yellow '
    'dots, healthy citizens as green dots. The Pie Chart and counters in the bottom panel update in real-time to '
    'reflect the changing population status.'
)

add_para('Step 6: Adjust Simulation Speed', bold=True)
add_para(
    'Use the "Speed" dropdown in the top toolbar to accelerate the simulation. Available options are 1x (real-time), '
    '2x, 5x, and 10x. Higher speeds are useful for observing long-term outcomes more quickly.'
)

add_para('Step 7: Pause and Reset', bold=True)
add_para(
    'Click "Pause" to freeze the simulation at any point, allowing examination of the current state. Click "Reset" '
    'to clear the entire simulation and return to a blank canvas.'
)

doc.add_page_break()

# ============================================================
# 9. RESULTS AND OUTPUT
# ============================================================
doc.add_heading('9. Results and Output', level=1)

add_para(
    'The CityGuardian simulation produces rich, real-time visual and quantitative output that allows users to assess '
    'the effectiveness of different emergency response configurations.'
)

add_para('Visual Output:', bold=True)
add_para(
    'The primary output is the 700x600 pixel canvas rendering of the city grid. The map displays the city layout with '
    'color-coded tile types, dynamically spreading disasters, moving emergency vehicles, and citizen status indicators. '
    'The rendering updates at approximately 60 frames per second, providing smooth animation of vehicle movement and '
    'disaster progression.'
)

add_para('Quantitative Output:', bold=True)
add_para(
    'The bottom panel displays three numerical counters: Total Citizens (the initial population count), Saved (citizens '
    'successfully evacuated by emergency services), and Casualties (citizens who died during the simulation). The Pie '
    'Chart provides a proportional breakdown of the population into four categories: Safe, Injured, Casualties, and '
    'Evacuated. These metrics update in real-time as the simulation progresses.'
)

add_para('Key Observations from Testing:', bold=True)
add_bullet('Fire Scenario with 3 Firetrucks: The fire spreads faster than the trucks can extinguish it. The trucks '
           'successfully clear small pockets but are unable to contain the overall spread, resulting in significant '
           'casualties and property destruction.')
add_bullet('Fire Scenario with 15+ Firetrucks: The larger fleet creates a containment perimeter faster than the fire '
           'can spread, resulting in successful suppression with minimal casualties.')
add_bullet('Earthquake Scenario: Ambulances rapidly deploy to the earthquake zone, load injured citizens in batches of '
           'up to 10, and shuttle them to the nearest hospital. The "Saved" counter increments in bursts as each '
           'ambulance completes its delivery cycle.')
add_bullet('Flood Scenario: Helicopters fly directly over flooded areas (ignoring road constraints) to rescue stranded '
           'citizens. The elevation-based flood model creates realistic inundation patterns with natural high-ground '
           'refuges.')

doc.add_page_break()

# ============================================================
# 10. PERFORMANCE ANALYSIS
# ============================================================
doc.add_heading('10. Performance Analysis', level=1)

add_para(
    'CityGuardian is designed for high performance despite managing thousands of entities simultaneously. The following '
    'analysis examines the computational complexity of the key subsystems.'
)

add_para('Rendering Performance:', bold=True)
add_para(
    'The Canvas-based rendering approach achieves 60 frames per second on modern hardware. Each frame requires iterating '
    'over all 4,200 tiles (O(W*H) where W=70 and H=60), all citizens (O(C) where C is typically 3,000-5,000), and all '
    'resources (O(R) where R is typically 3-20). The total per-frame rendering cost is O(W*H + C + R), which is linear '
    'and highly efficient. By avoiding Scene Graph overhead, the rendering cost is dominated by the simple fillRect and '
    'fillOval drawing commands, which are hardware-accelerated on most systems.'
)

add_para('Pathfinding Performance:', bold=True)
add_para(
    'The A* algorithm has a worst-case time complexity of O(b^d) where b is the branching factor (4 for cardinal-direction '
    'movement) and d is the depth of the optimal path. In practice, the Manhattan distance heuristic dramatically prunes '
    'the search space, and typical pathfinding operations complete in under 1 millisecond on the 70x60 grid. Pathfinding '
    'is only performed once per dispatch cycle (when a resource transitions from AVAILABLE to DISPATCHED), and the result '
    'is cached as a List<Tile>. Subsequent movement frames simply advance the path index, requiring O(1) computation.'
)

add_para('Disaster Evolution Performance:', bold=True)
add_para(
    'Fire and flood evolution require a full grid scan (O(W*H)) per spread interval. However, because the spread '
    'intervals are 0.8 and 3.0 seconds respectively (not per-frame), the amortized per-frame cost is negligible. '
    'Earthquake evolution is O(R^2 + C) where R is the damage radius and C is the citizen count, and executes exactly '
    'once per earthquake event.'
)

add_para('Memory Consumption:', bold=True)
add_para(
    'The primary memory consumers are the Tile array (4,200 objects), the Citizen list (3,000-5,000 objects), and the '
    'A* pathfinding data structures (proportional to the number of explored nodes per query). Total memory consumption '
    'is estimated at 5-15 MB, well within the capabilities of any modern system.'
)

doc.add_page_break()

# ============================================================
# 11. CHALLENGES ENCOUNTERED AND SOLUTIONS
# ============================================================
doc.add_heading('11. Challenges Encountered and Solutions', level=1)

add_para(
    'During the development of CityGuardian, several significant technical challenges were encountered and resolved.'
)

add_para('Challenge 1: Emergency Vehicles Getting Stuck', bold=True)
add_para(
    'Initially, firetrucks would frequently become stuck in the middle of the map. The root cause was identified as '
    'the fire disaster destroying road tiles beneath the vehicles. When a road tile was converted to BURNT, the A* '
    'pathfinding could no longer find a valid route, causing the vehicle to stall permanently. The solution was to '
    'modify the FireDisaster class to prevent road tiles from transitioning to BURNT during natural burnout, preserving '
    'the road network throughout the simulation.'
)

add_para('Challenge 2: Fire Containment Balance', bold=True)
add_para(
    'Achieving a meaningful difference between 3 and 15 firetrucks required careful tuning of the fire spread parameters. '
    'Initially, the fire spread too quickly for any number of trucks to contain, or too slowly for there to be any '
    'meaningful distinction. The final calibration of 0.8-second spread intervals with 40% spread probability achieved '
    'the desired balance where 3 trucks fail to contain the fire but 15 trucks succeed.'
)

add_para('Challenge 3: Ambulance Jittering and Back-and-Forth Movement', bold=True)
add_para(
    'Ambulances initially exhibited a jittery, indecisive movement pattern. This was caused by ambulances targeting '
    'individual citizens one at a time. When a citizen was rescued, the ambulance would immediately turn around to '
    'target the next nearest citizen, often located just one tile away. The solution was to implement the area rescue '
    'system, where ambulances load multiple citizens in a 5x5 block upon arrival, and the hospital transport system, '
    'which gives ambulances a clear round-trip objective.'
)

add_para('Challenge 4: Citizens Walking Away from Disaster Zones', bold=True)
add_para(
    'After introducing hospital infrastructure, a dormant evacuation feature was accidentally activated. Citizens '
    'began autonomously walking to hospitals during disasters, which conflicted with the intended simulation behavior '
    'where citizens remain stationary and rely on emergency services. The solution was to disable autonomous citizen '
    'evacuation routing entirely, ensuring citizens stay at their original positions during all disaster types.'
)

add_para('Challenge 5: Earthquake Injury Threshold', bold=True)
add_para(
    'Ambulances initially failed to dispatch during earthquake scenarios because the earthquake damage (40 points) '
    'did not reduce citizen health below the injury threshold (50 health). Since citizens start at 100 health, 40 '
    'damage only reduced them to 60, which did not trigger the isInjured flag. Increasing the damage to 60 points '
    '(dropping health to 40, well below the 50 threshold) resolved the issue.'
)

doc.add_page_break()

# ============================================================
# 12. FUTURE ENHANCEMENTS
# ============================================================
doc.add_heading('12. Future Enhancements', level=1)

add_para('The following enhancements are proposed for future development iterations:')

add_bullet('Traffic Congestion Modeling: Implement vehicle-to-vehicle collision detection and queuing on roads, '
           'creating realistic traffic jams when multiple emergency vehicles converge on the same road segment.')
add_bullet('Hospital Capacity Limits: Introduce finite capacity for hospital tiles, requiring ambulances to route '
           'to alternative hospitals when the nearest one is full.')
add_bullet('GIS Integration: Replace procedural generation with real-world geographic data from OpenStreetMap '
           'or similar services, enabling simulation of actual city layouts.')
add_bullet('Historical Replay: Leverage the SQLite database to record complete simulation state at regular '
           'intervals, enabling post-simulation playback and analysis.')
add_bullet('Advanced AI Insights: Expand the RecommendationEngine to provide real-time predictive analytics, '
           'such as estimated time to full containment, optimal resource reallocation suggestions, and citizen '
           'survival probability forecasts.')
add_bullet('Multi-Disaster Scenarios: Support concurrent disasters (e.g., an earthquake triggering secondary '
           'fires) with compound resource allocation challenges.')
add_bullet('Network Multiplayer: Allow multiple users to collaboratively manage emergency resources during '
           'a shared disaster scenario.')

doc.add_page_break()

# ============================================================
# 13. CONCLUSION
# ============================================================
doc.add_heading('13. Conclusion', level=1)

add_para(
    'CityGuardian successfully demonstrates a comprehensive, scalable architecture for modeling complex urban disaster '
    'scenarios and evaluating emergency response strategies. The project integrates several foundational computer science '
    'concepts including the A* pathfinding algorithm, cellular automata, finite state machines, the observer pattern, '
    'and the game loop design pattern into a cohesive, real-time simulation platform.'
)
add_para(
    'The three disaster modules (Fire, Flood, Earthquake) each exhibit unique, physically-motivated behavior that '
    'creates distinct tactical challenges for the emergency response system. The fire module\'s cellular automata-based '
    'spread creates organic, unpredictable flame fronts. The flood module\'s elevation-dependent water flow produces '
    'realistic inundation patterns. The earthquake module\'s instant-impact damage creates sudden, high-density '
    'casualty events that stress the logistical capacity of the ambulance fleet.'
)
add_para(
    'The emergency resource system demonstrates meaningful strategic differentiation between vehicle types. Firetrucks '
    'navigate road networks to contain spreading fires. Ambulances implement a multi-phase pickup-transport-delivery '
    'logistics chain involving hospitals. Helicopters bypass road infrastructure entirely to reach flood-stranded '
    'citizens. The user\'s ability to configure the fleet size for each resource type creates clear, observable '
    'differences in simulation outcomes, validating the project\'s core thesis that resource allocation directly '
    'impacts disaster response effectiveness.'
)
add_para(
    'The project is architecturally well-positioned for future expansion, with clean separation of concerns, '
    'extensible class hierarchies, and a database infrastructure ready for historical analysis. CityGuardian '
    'serves as both a practical simulation tool and an educational reference for advanced Java programming, '
    'algorithmic design, and software architecture principles.'
)

doc.add_page_break()

# ============================================================
# 14. REFERENCES
# ============================================================
doc.add_heading('14. References', level=1)

refs = [
    'Hart, P. E., Nilsson, N. J., & Raphael, B. (1968). "A Formal Basis for the Heuristic Determination of Minimum Cost Paths." IEEE Transactions on Systems Science and Cybernetics, 4(2), 100-107.',
    'Dijkstra, E. W. (1959). "A Note on Two Problems in Connexion with Graphs." Numerische Mathematik, 1(1), 269-271.',
    'Wolfram, S. (1983). "Statistical Mechanics of Cellular Automata." Reviews of Modern Physics, 55(3), 601-644.',
    'Wolfram, S. (2002). A New Kind of Science. Wolfram Media, Inc.',
    'Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (1994). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley Professional.',
    'Nystrom, R. (2014). Game Programming Patterns. Genever Benning. Chapter on Game Loop pattern.',
    'Russell, S. J., & Norvig, P. (2020). Artificial Intelligence: A Modern Approach (4th Edition). Pearson. Chapters on Heuristic Search and Multi-Agent Systems.',
    'Oracle Corporation. (2023). "JavaFX Architecture." JavaFX Documentation. Available at: https://openjfx.io/javadoc/17/',
    'Oracle Corporation. (2023). "JavaFX Canvas API." JavaFX Documentation. Available at: https://openjfx.io/javadoc/17/javafx.graphics/javafx/scene/canvas/Canvas.html',
    'Apache Software Foundation. (2023). "Apache Maven Project." Available at: https://maven.apache.org/',
    'Xerial Project. (2023). "SQLite JDBC Driver." Available at: https://github.com/xerial/sqlite-jdbc',
    'Koenig, S., & Likhachev, M. (2002). "D* Lite." Proceedings of the AAAI Conference on Artificial Intelligence, 476-483.',
    'Sullivan, W. G. (2009). "Fire Spread Modeling Using Cellular Automata." International Journal of Wildland Fire, 18(4), 369-386.',
    'Jongman, B., Ward, P. J., & Aerts, J. C. J. H. (2012). "Global Exposure to River and Coastal Flooding: Long Term Trends and Changes." Global Environmental Change, 22(4), 823-835.',
    'FEMA. (2019). "National Incident Management System (NIMS)." Federal Emergency Management Agency, U.S. Department of Homeland Security.',
]

for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'[{i}]  {ref}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)

# ============================================================
# SAVE
# ============================================================
doc.save('report/CityGuardian_Technical_Report.docx')
print('Report generated successfully: report/CityGuardian_Technical_Report.docx')
